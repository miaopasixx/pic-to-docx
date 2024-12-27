from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import tempfile
from PIL import Image, ImageDraw, ImageFont
import io
import base64

def create_preview(title, img1_path, img2_path):
    """生成预览图像的base64字符串"""
    try:
        # 创建一个A4大小的白色图像（300dpi）
        width = 2480
        height = 3508
        image = Image.new('RGB', (width, height), 'white')
        draw = ImageDraw.Draw(image)
        
        # 尝试加载字体，如果失败则使用默认字体
        try:
            font = ImageFont.truetype('simhei.ttf', 120)  # 大约对应22pt
        except:
            font = ImageFont.load_default()
            
        # 绘制标题
        title_bbox = draw.textbbox((0, 0), title, font=font)
        title_width = title_bbox[2] - title_bbox[0]
        title_x = (width - title_width) // 2
        draw.text((title_x, 100), title, font=font, fill='black')
        
        # 计算图片区域
        available_width = int(width * 0.9)  # 留出10%边距
        available_height = int((height - 300) // 2)  # 减去标题区域后平分
        
        # 处理第一张图片
        img1 = Image.open(img1_path)
        img1 = resize_image(img1, available_width, available_height)
        x1 = (width - img1.width) // 2
        y1 = 300
        image.paste(img1, (x1, y1))
        
        # 处理第二张图片
        img2 = Image.open(img2_path)
        img2 = resize_image(img2, available_width, available_height)
        x2 = (width - img2.width) // 2
        y2 = y1 + available_height + 100
        image.paste(img2, (x2, y2))
        
        # 转换为base64
        buffer = io.BytesIO()
        image.save(buffer, format='PNG')
        base64_str = base64.b64encode(buffer.getvalue()).decode()
        
        return base64_str
        
    except Exception as e:
        print(f"Error in create_preview: {str(e)}")
        raise Exception(f"预览生成失败: {str(e)}")

def resize_image(img, max_width, max_height):
    """保持宽高比例调整图片大小"""
    ratio = min(max_width/img.width, max_height/img.height)
    new_width = int(img.width * ratio)
    new_height = int(img.height * ratio)
    return img.resize((new_width, new_height), Image.Resampling.LANCZOS)

def create_photo_doc(title, img1_path, img2_path, output_path=None):
    try:
        print(f"Creating document with title: {title}")
        print(f"Image paths: {img1_path}, {img2_path}")
        
        # 创建文档对象
        doc = Document()
        
        # 设置标题
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 设置标题字体
        try:
            title_run.font.name = '方正小标宋简体'
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
        except:
            title_run.font.name = 'SimSun'
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            
        title_run.font.size = Pt(22)
        
        # 添加标题下方的空格
        space_paragraph = doc.add_paragraph()
        space_run = space_paragraph.add_run("\n")
        space_run.font.size = Pt(2)  # 三号字体大约为16pt
        
        # A4纸张设置
        section = doc.sections[0]
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.left_margin = Inches(0.25)  # 减小边距
        section.right_margin = Inches(0.25)
        section.top_margin = Inches(0.25)
        section.bottom_margin = Inches(0.25)
        
        # 计算图片尺寸
        available_width = Inches(7.27)  # 8.27 - 0.5(margins)
        available_height = Inches(4.5)   # (11.69 - 0.5(margins) - 0.5(title)) / 2
        
        # 添加图片1
        if os.path.exists(img1_path):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(img1_path, width=available_width, height=available_height)
        else:
            raise Exception(f"找不到第一张图片: {img1_path}")
        
        # 添加图片之间的空格
        space_paragraph = doc.add_paragraph()
        space_run = space_paragraph.add_run("\n")
        space_run.font.size = Pt(16)  # 三号字体大约为16pt
        
        # 添加图片2
        if os.path.exists(img2_path):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(img2_path, width=available_width, height=available_height)
        else:
            raise Exception(f"找不到第二张图片: {img2_path}")
        
        # 保存文档
        if output_path is None:
            output_path = f"照片丨{title}.docx"
        doc.save(output_path)
        print(f"Document saved successfully to: {output_path}")
        
        return output_path
        
    except Exception as e:
        print(f"Error in create_photo_doc: {str(e)}")
        raise Exception(f"文档生成失败: {str(e)}")