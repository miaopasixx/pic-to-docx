import tkinter as tk
from tkinter import messagebox
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import fnmatch

# 显示使用说明
def show_usage_guide():
    usage = """照片合并工具使用说明：

1. 将程序放在需要处理的目标文件夹中
2. 确保目标文件夹下有一个名为"照片丨xxx"的子文件夹
3. 子文件夹中需包含两张照片，命名为"1.jpg"和"2.jpg"
4. 运行程序后将自动生成Word文档
5. 生成的文档名为"照片丨当前文件夹名称.docx"

文档格式：A4纸张，标题使用方正小标宋简体"""
    
    root = tk.Tk()
    root.withdraw()  # 隐藏主窗口
    messagebox.showinfo("使用说明", usage)
    root.destroy()

def is_first_run():
    flag_file = os.path.join(os.path.expanduser("~"), ".photo_combiner_flag")
    if not os.path.exists(flag_file):
        # 创建标记文件
        with open(flag_file, "w") as f:
            f.write("1")
        return True
    return False

def create_photo_doc():
    # 检查是否首次运行
    if is_first_run():
        show_usage_guide()
    
    try:
        # 创建文档对象
        doc = Document()
        
        # 获取当前文件夹名称
        current_folder_name = os.path.basename(os.getcwd())
        
        # 查找包含"照片丨"的文件夹
        folder_name = None
        for name in os.listdir(os.getcwd()):
            if fnmatch.fnmatch(name, '*照片丨*'):
                folder_name = name
                break

        if folder_name is None:
            raise FileNotFoundError("未找到包含'照片丨'的文件夹")

        folder_path = os.path.join(os.getcwd(), folder_name)
        
        # 设置标题为当前文件夹的名称
        title = doc.add_paragraph()
        title_run = title.add_run(current_folder_name)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 设置标题字体为方正小标宋简体，二号字体
        title_run.font.name = '方正小标宋简体'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
        title_run.font.size = Pt(22)  # 修改为二号字体大小（22磅）
        
        # 修改A4纸张计算部分
        page_width = 8.27  # A4纸张宽度（英寸）
        page_height = 11.69  # A4纸张高度（英寸）
        margin = 0.5  # 进一步减小页边距
        available_width = page_width - 2 * margin
        
        # 进一步压缩空间分配
        title_height = 0.2  # 减小标题高度
        spacing = 0.1  # 减小图片间距
        available_height = (page_height - 2 * margin - title_height - spacing) / 2.2  # 稍微减小图片高度
        
        # 设置段落和文档属性
        section = doc.sections[0]
        section.page_height = Inches(page_height)
        section.page_width = Inches(page_width)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)
        
        # 设置标题段落属性
        title.space_after = Pt(5)
        title.space_before = Pt(5)
        
        # 添加图片1
        img1_path = os.path.join(folder_path, "1.jpg")
        if os.path.exists(img1_path):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.space_before = Pt(0)
            paragraph.space_after = Pt(0)
            run = paragraph.add_run()
            run.add_picture(img1_path, width=Inches(available_width), height=Inches(available_height))
        
        # 添加图片2
        img2_path = os.path.join(folder_path, "2.jpg")
        if os.path.exists(img2_path):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.space_before = Pt(0)
            paragraph.space_after = Pt(0)
            run = paragraph.add_run()
            run.add_picture(img2_path, width=Inches(available_width), height=Inches(available_height))
        
        # 保存文档，名称为"照片丨当前文件夹的名称"
        doc_name = f"照片丨{current_folder_name}.docx"
        doc.save(doc_name)
    except FileNotFoundError as e:
        root = tk.Tk()
        root.withdraw()
        messagebox.showerror("错误", str(e))
        root.destroy()
        return
    except Exception as e:
        print(f"Error: {e}")  # 打印错误信息到控制台
        return str(e), 500

if __name__ == "__main__":
    create_photo_doc()